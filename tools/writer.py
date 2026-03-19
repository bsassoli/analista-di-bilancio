"""Writer per output Excel (.xlsx) e Word (.docx)."""

import re
from pathlib import Path
from typing import Any, Optional

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# --- Formattazione numeri italiani ---

def formatta_numero_it(valore: int | float | None) -> str:
    """Formatta un numero in formato italiano (punto migliaia, niente decimali).

    Es: 1250000 → "1.250.000", -350000 → "(350.000)"
    """
    if valore is None:
        return "n/d"
    if valore == 0:
        return "0"

    negativo = valore < 0
    v = abs(int(valore))
    formatted = f"{v:,}".replace(",", ".")

    if negativo:
        return f"({formatted})"
    return formatted


def formatta_percentuale(valore: float | None, decimali: int = 1) -> str:
    """Formatta un valore come percentuale italiana.

    Es: 0.123 → "12,3%"
    """
    if valore is None:
        return "n/d"
    pct = valore * 100
    formatted = f"{pct:.{decimali}f}".replace(".", ",")
    return f"{formatted}%"


def formatta_indice(valore: float | None, decimali: int = 2) -> str:
    """Formatta un indice con virgola decimale.

    Es: 0.85 → "0,85"
    """
    if valore is None:
        return "n/d"
    return f"{valore:.{decimali}f}".replace(".", ",")


# --- Excel writer ---

# Stili riutilizzabili
FONT_HEADER = Font(bold=True, size=11)
FONT_TOTALE = Font(bold=True, size=10)
FONT_NORMALE = Font(size=10)
FONT_NEGATIVE = Font(size=10, color="FF0000")
FILL_HEADER = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
FONT_HEADER_WHITE = Font(bold=True, size=11, color="FFFFFF")
FILL_GREEN = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
FILL_YELLOW = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
FILL_RED = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
BORDER_BOTTOM = Border(bottom=Side(style="thin"))


def scrivi_excel(path: str, fogli: dict[str, list[list[Any]]]) -> str:
    """Scrive un file Excel con più fogli.

    Args:
        path: Path di output.
        fogli: Dict nome_foglio → lista di righe (ogni riga è una lista di valori).

    Returns:
        Path del file scritto.
    """
    wb = Workbook()

    for i, (nome, righe) in enumerate(fogli.items()):
        if i == 0:
            ws = wb.active
            ws.title = nome
        else:
            ws = wb.create_sheet(title=nome)

        for r_idx, riga in enumerate(righe, start=1):
            for c_idx, valore in enumerate(riga, start=1):
                cell = ws.cell(row=r_idx, column=c_idx, value=valore)

                # Intestazione (prima riga)
                if r_idx == 1:
                    cell.font = FONT_HEADER_WHITE
                    cell.fill = FILL_HEADER
                    cell.alignment = Alignment(horizontal="center")

        # Auto-width colonne
        for col in range(1, ws.max_column + 1):
            max_len = 0
            for row in range(1, ws.max_row + 1):
                val = ws.cell(row=row, column=col).value
                if val:
                    max_len = max(max_len, len(str(val)))
            ws.column_dimensions[get_column_letter(col)].width = min(max_len + 4, 40)

    out_path = Path(path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out_path))
    return str(out_path)


def crea_tabella_serie_storica(
    dati: dict[str, dict[str, Any]], anni: list[str], label_col: str = "Voce"
) -> list[list[Any]]:
    """Prepara dati per una tabella serie storica multi-anno.

    Args:
        dati: Dict voce_label → {anno: valore}.
        anni: Lista anni ordinati.
        label_col: Intestazione colonna label.

    Returns:
        Lista di righe (prima riga = intestazione).
    """
    righe = [[label_col] + anni]
    for label, valori_per_anno in dati.items():
        riga = [label]
        for anno in anni:
            val = valori_per_anno.get(anno)
            riga.append(formatta_numero_it(val) if isinstance(val, (int, float)) else val)
        righe.append(riga)
    return righe


# --- Word writer ---

def scrivi_word(path: str, sezioni: list[dict]) -> str:
    """Scrive un documento Word con sezioni strutturate.

    Args:
        path: Path di output.
        sezioni: Lista di dict con chiavi:
            - tipo: "titolo" | "sottotitolo" | "paragrafo" | "tabella" | "lista"
            - contenuto: str (per testo) o list[list[str]] (per tabella)
            - livello: int (per titoli, default 1)

    Returns:
        Path del file scritto.
    """
    doc = Document()

    # Stile base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    for sezione in sezioni:
        tipo = sezione.get("tipo", "paragrafo")
        contenuto = sezione.get("contenuto", "")
        livello = sezione.get("livello", 1)

        if tipo == "titolo":
            doc.add_heading(contenuto, level=livello)

        elif tipo == "sottotitolo":
            doc.add_heading(contenuto, level=livello + 1)

        elif tipo == "paragrafo":
            doc.add_paragraph(contenuto)

        elif tipo == "tabella":
            if not contenuto or not isinstance(contenuto, list):
                continue
            table = doc.add_table(rows=len(contenuto), cols=len(contenuto[0]))
            table.style = "Light Grid Accent 1"
            for r_idx, riga in enumerate(contenuto):
                for c_idx, val in enumerate(riga):
                    cell = table.cell(r_idx, c_idx)
                    cell.text = str(val) if val is not None else ""
                    if r_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

        elif tipo == "lista":
            if isinstance(contenuto, list):
                for item in contenuto:
                    doc.add_paragraph(str(item), style="List Bullet")

    out_path = Path(path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return str(out_path)
